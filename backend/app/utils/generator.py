import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document

from backend.app.config import GENERATED_DIR
from backend.app.database import get_connection
from backend.app.utils.manual_docx_filler import apply_manual_fill_to_generated_docx
from backend.app.utils.teaching_load import (
    build_teaching_load_context,
    build_teaching_load_summary,
    is_teaching_load_summary_raw_table,
)


TOTAL_TEXT_VARIANTS = ("итого", "итог", "total", "всего", "барлығы")
NUMERIC_TOTAL_FIELDS = (
    "l",
    "spz",
    "lz",
    "srsp",
    "rk_1_2",
    "ekzameny",
    "practika",
    "diploma_supervision",
    "research_work",
    "other_work",
    "itogo",
)
PAYLOAD_FIELDS = (
    "discipline",
    "op",
    "group",
    "course",
    "academic_period",
    "credits",
    "student_count",
    "l",
    "spz",
    "lz",
    "srsp",
    "rk_1_2",
    "ekzameny",
    "practika",
    "diploma_supervision",
    "research_work",
    "other_work",
    "itogo",
)
SUMMARY_TABLE_COLUMN_MAP = {
    "l": 1,
    "spz": 2,
    "lz": 3,
    "srsp": 4,
    "rk_1_2": 5,
    "ekzameny": 6,
    "class_hours": 7,
    "practika": 8,
    "research_work": 9,
    "diploma_supervision": 10,
    "other_work": 11,
    "office_hours": 12,
    "itogo": 13,
}


def _normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _to_str(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.2f}".rstrip("0").rstrip(".")
    return str(value)


def _safe_name(value: str) -> str:
    return re.sub(r"[^\w]+", "_", str(value or ""), flags=re.UNICODE).strip("_")


def _safe_get_table(doc: Document, table_index: int):
    if table_index < 0 or table_index >= len(doc.tables):
        return None
    return doc.tables[table_index]


def _safe_get_cell(table, row_index: int, col_index: int):
    try:
        return table.rows[row_index].cells[col_index]
    except Exception:
        return None


def _set_cell_text(cell, value: Any):
    if cell is None:
        return
    cell.text = _to_str(value)


def _clear_row(row):
    for cell in row.cells:
        cell.text = ""


def _clone_row_before(table, row_index: int, source_row_index: Optional[int] = None):
    source_index = row_index if source_row_index is None else source_row_index
    tr = table.rows[source_index]._tr
    new_tr = deepcopy(tr)
    table.rows[row_index]._tr.addprevious(new_tr)
    return row_index


def _extract_teacher(cur, teacher_id: int) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT
            t.id,
            t.full_name,
            t.department_id,
            t.faculty,
            t.position,
            t.academic_degree,
            t.academic_rank,
            t.staff_type,
            t.extra_data,
            d.name
        FROM teachers t
        LEFT JOIN departments d ON d.id = t.department_id
        WHERE t.id = %s;
        """,
        (teacher_id,),
    )
    row = cur.fetchone()
    if not row:
        raise Exception("Преподаватель не найден")

    return {
        "id": row[0],
        "full_name": row[1] or "",
        "department_id": row[2],
        "faculty": row[3] or "",
        "position": row[4] or "",
        "academic_degree": row[5] or "",
        "academic_rank": row[6] or "",
        "staff_type": row[7] or "",
        "extra_data": row[8] or {},
        "department": row[9] or "",
    }


def _get_excel_by_year(cur, department_id: int, academic_year: str) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT id
        FROM excel_templates
        WHERE department_id = %s AND academic_year = %s;
        """,
        (department_id, academic_year),
    )
    row = cur.fetchone()
    if not row:
        raise Exception("Для этого года Excel не загружен")
    return {"id": int(row[0])}


def _get_raw_template_by_year(cur, department_id: int, academic_year: str) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT id, file_path
        FROM raw_docx_templates
        WHERE department_id = %s AND academic_year = %s;
        """,
        (department_id, academic_year),
    )
    row = cur.fetchone()
    if not row:
        raise Exception("Для этого года raw шаблон не загружен")
    return {"id": int(row[0]), "file_path": row[1]}


def _get_settings_for_excel(cur, excel_template_id: int) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT config
        FROM generation_settings
        WHERE excel_template_id = %s
        LIMIT 1;
        """,
        (excel_template_id,),
    )
    row = cur.fetchone()
    if not row:
        raise Exception("Нет настроек для этого Excel. Сначала сохрани Settings.")
    return row[0] or {}


def _get_excel_columns(cur, excel_template_id: int) -> List[Tuple[str, str]]:
    cur.execute(
        """
        SELECT column_name, header_text
        FROM excel_columns
        WHERE template_id = %s
        ORDER BY position_index;
        """,
        (excel_template_id,),
    )
    return cur.fetchall()


def _get_excel_rows(cur, excel_template_id: int) -> List[Dict[str, Any]]:
    cur.execute(
        """
        SELECT row_data
        FROM excel_rows
        WHERE template_id = %s
        ORDER BY row_number;
        """,
        (excel_template_id,),
    )
    return [row[0] for row in cur.fetchall()]


def _get_raw_tables(cur, raw_template_id: int) -> Dict[int, Dict[str, Any]]:
    cur.execute(
        """
        SELECT
            id,
            table_index,
            section_title,
            table_type,
            row_count,
            col_count,
            header_signature,
            has_total_row,
            loop_template_row_index,
            column_hints,
            editable_cells_count,
            prefilled_cells_count,
            table_fingerprint,
            structure_meta,
            extra_meta
        FROM raw_docx_tables
        WHERE template_id = %s
        ORDER BY table_index;
        """,
        (raw_template_id,),
    )
    out: Dict[int, Dict[str, Any]] = {}
    for row in cur.fetchall():
        out[int(row[0])] = {
            "id": int(row[0]),
            "table_index": int(row[1]),
            "section_title": row[2] or "",
            "table_type": row[3] or "",
            "row_count": int(row[4] or 0),
            "col_count": int(row[5] or 0),
            "header_signature": row[6] or "",
            "has_total_row": bool(row[7]),
            "loop_template_row_index": row[8] if row[8] is None else int(row[8]),
            "column_hints": row[9] or [],
            "editable_cells_count": int(row[10] or 0),
            "prefilled_cells_count": int(row[11] or 0),
            "table_fingerprint": row[12] or "",
            "structure_meta": row[13] or {},
            "extra_meta": row[14] or {},
        }
    return out


def _build_excel_context(
    teacher: Dict[str, Any],
    excel_columns: List[Tuple[str, str]],
    excel_rows: List[Dict[str, Any]],
    settings_cfg: Dict[str, Any],
    academic_year: str,
) -> Dict[str, Any]:
    return build_teaching_load_context(
        teacher=teacher,
        excel_columns=excel_columns,
        excel_rows=excel_rows,
        settings_cfg=settings_cfg,
        academic_year=academic_year,
    )


def _row_text(table_row) -> str:
    return " ".join(_normalize_text(cell.text).lower() for cell in table_row.cells if _normalize_text(cell.text)).strip()


def _is_total_text(text: str) -> bool:
    return any(variant in text for variant in TOTAL_TEXT_VARIANTS)


def _parse_scope_from_text(text: str) -> tuple[int, ...]:
    norm = _normalize_text(text).lower()
    if not norm:
        return ()
    if "контроль" in norm and not _is_total_text(norm):
        return ()

    match = re.match(r"^\s*(\d+(?:\s*,\s*\d+)*)", norm)
    if not match:
        return ()

    if "сем" not in norm and "sem" not in norm and not re.fullmatch(r"\d+(?:\s*,\s*\d+)*", norm):
        return ()

    numbers: List[int] = []
    for part in match.group(1).split(","):
        try:
            number = int(part.strip())
        except Exception:
            continue
        if 0 < number <= 12 and number not in numbers:
            numbers.append(number)
    return tuple(numbers)


def _detect_table_blocks(table, raw_table: Dict[str, Any]) -> List[Dict[str, Any]]:
    default_insert_start = int(raw_table.get("loop_template_row_index") or 1)
    rows_meta: List[Dict[str, Any]] = []

    for row_index, row in enumerate(table.rows):
        text = _row_text(row)
        scope = _parse_scope_from_text(text)
        rows_meta.append(
            {
                "row_index": row_index,
                "scope": scope,
                "is_total": bool(scope and _is_total_text(text)),
            }
        )

    total_rows = [item for item in rows_meta if item["is_total"]]
    blocks: List[Dict[str, Any]] = []
    prev_total_row_index: Optional[int] = None

    for total_row in total_rows:
        search_start = default_insert_start if prev_total_row_index is None else prev_total_row_index + 1
        label_row_index = None

        for candidate in rows_meta:
            candidate_row_index = int(candidate["row_index"])
            if candidate_row_index < search_start or candidate_row_index >= int(total_row["row_index"]):
                continue
            if candidate["scope"] == total_row["scope"] and not candidate["is_total"]:
                label_row_index = candidate_row_index

        insert_start = label_row_index + 1 if label_row_index is not None else search_start
        if insert_start > int(total_row["row_index"]):
            insert_start = int(total_row["row_index"])

        blocks.append(
            {
                "scope": tuple(total_row["scope"]),
                "scope_key": ",".join(str(x) for x in total_row["scope"]),
                "label_row_index": label_row_index,
                "insert_start_row_index": insert_start,
                "total_row_index": int(total_row["row_index"]),
            }
        )
        prev_total_row_index = int(total_row["row_index"])

    return blocks


def _guess_column_map(raw_table: Dict[str, Any]) -> Dict[str, int]:
    hints = [str(value).strip().lower() for value in (raw_table.get("column_hints") or [])]
    out: Dict[str, int] = {}

    for idx, hint in enumerate(hints):
        if "наименование" in hint or "subject" in hint or "пән" in hint:
            out["discipline"] = idx
        elif "образовательная программа" in hint or hint == "оп" or "program" in hint:
            out["op"] = idx
        elif "группа" in hint or "group" in hint:
            out["group"] = idx
        elif "академ" in hint or "period" in hint:
            out["academic_period"] = idx
        elif "курс" in hint or hint == "course":
            out["course"] = idx
        elif "кредит" in hint:
            out["credits"] = idx
        elif "обуча" in hint or "контингент" in hint or "students" in hint:
            out["student_count"] = idx
        elif "лек" in hint:
            out["l"] = idx
        elif "практ" in hint:
            out["spz"] = idx
        elif "лабор" in hint:
            out["lz"] = idx
        elif "срсп" in hint or "сроп" in hint:
            out["srsp"] = idx
        elif "рубеж" in hint:
            out["rk_1_2"] = idx
        elif "экзам" in hint:
            out["ekzameny"] = idx
        elif "практика" in hint:
            out["practika"] = idx
        elif "рук-во дп" in hint or "дп и мд" in hint or "диссертац" in hint:
            out["diploma_supervision"] = idx
        elif "нирм" in hint or "нирд" in hint:
            out["research_work"] = idx
        elif "двр" in hint or "другой" in hint or "дополнительн" in hint:
            out["other_work"] = idx
        elif "итого" in hint and "час" in hint:
            out["itogo"] = idx

    fallback_indexes = {
        "discipline": 1,
        "group": 2,
    }

    for field_key, fallback_index in fallback_indexes.items():
        if field_key in out:
            continue
        if len(hints) > fallback_index:
            out[field_key] = fallback_index

    return out


def _display_value(value: Any) -> Any:
    if value is None:
        return ""
    if isinstance(value, (int, float)) and abs(float(value)) < 1e-9:
        return ""
    return value


def _build_payload(row_data: Dict[str, Any]) -> Dict[str, Any]:
    return {
        field_key: _display_value(row_data.get(field_key, ""))
        for field_key in PAYLOAD_FIELDS
    }


def _fill_row_by_map(table, row_index: int, payload: Dict[str, Any], col_map: Dict[str, int]):
    for field_key, col_index in (col_map or {}).items():
        if field_key not in payload:
            continue
        cell = _safe_get_cell(table, row_index, int(col_index))
        _set_cell_text(cell, payload.get(field_key))


def _fill_total_row(table, row_index: int, totals: Dict[str, Any], col_map: Dict[str, int]):
    payload = {
        field_key: _display_value(value)
        for field_key, value in (totals or {}).items()
    }
    _fill_row_by_map(
        table=table,
        row_index=row_index,
        payload={field_key: payload.get(field_key, "") for field_key in NUMERIC_TOTAL_FIELDS},
        col_map=col_map,
    )


def _render_scope_block(
    table,
    block: Dict[str, Any],
    rows_data: List[Dict[str, Any]],
    totals: Dict[str, Any],
    col_map: Dict[str, int],
):
    insert_start_row_index = int(block["insert_start_row_index"])
    total_row_index = int(block["total_row_index"])
    available_slots = max(total_row_index - insert_start_row_index, 0)
    payloads = [_build_payload(row) for row in (rows_data or [])]

    if len(payloads) > available_slots:
        need_add = len(payloads) - available_slots
        for _ in range(need_add):
            source_row_index = max(insert_start_row_index, total_row_index - 1)
            _clone_row_before(table, total_row_index, source_row_index=source_row_index)
            total_row_index += 1

    for idx, payload in enumerate(payloads):
        row_index = insert_start_row_index + idx
        if row_index >= total_row_index:
            break
        _clear_row(table.rows[row_index])
        _fill_row_by_map(table, row_index, payload, col_map)

    start_clear = insert_start_row_index + len(payloads)
    for row_index in range(start_clear, total_row_index):
        _clear_row(table.rows[row_index])

    _fill_total_row(table, total_row_index, totals or {}, col_map)


def _find_annual_total_row_index(table, last_semester_total_row_index: int) -> Optional[int]:
    for row_index in range(last_semester_total_row_index + 1, len(table.rows)):
        text = _row_text(table.rows[row_index])
        if _is_total_text(text) and not _parse_scope_from_text(text):
            return row_index
    return None


def _map_scope_rows_to_blocks(
    rows_by_scope: Dict[str, List[Dict[str, Any]]],
    blocks: List[Dict[str, Any]],
    primary_common_scope_key: Optional[str],
) -> Dict[str, List[Dict[str, Any]]]:
    mapped: Dict[str, List[Dict[str, Any]]] = {
        block["scope_key"]: list(rows_by_scope.get(block["scope_key"]) or [])
        for block in blocks
    }
    available_scope_keys = {block["scope_key"] for block in blocks}
    common_block_keys = [block["scope_key"] for block in blocks if "," in block["scope_key"]]

    for scope_key, rows in (rows_by_scope or {}).items():
        if scope_key in available_scope_keys:
            continue

        fallback_scope_key = None
        if primary_common_scope_key and primary_common_scope_key in available_scope_keys:
            fallback_scope_key = primary_common_scope_key
        elif common_block_keys:
            fallback_scope_key = common_block_keys[-1]

        if not fallback_scope_key:
            continue

        mapped.setdefault(fallback_scope_key, []).extend(rows)

    for scope_key, rows in mapped.items():
        rows.sort(key=lambda item: item.get("_row_order") or 0)
        mapped[scope_key] = rows

    return mapped


def _sum_scope_rows(rows: List[Dict[str, Any]]) -> Dict[str, float]:
    totals = {field_key: 0.0 for field_key in NUMERIC_TOTAL_FIELDS}
    for row in rows or []:
        for field_key in NUMERIC_TOTAL_FIELDS:
            value = row.get(field_key)
            if value is None:
                continue
            try:
                totals[field_key] += float(value)
            except Exception:
                continue
    return totals


def _resolve_teaching_load_summary_raw_table(
    raw_tables: Dict[int, Dict[str, Any]],
    settings_cfg: Dict[str, Any],
) -> Optional[Dict[str, Any]]:
    template_bindings = (settings_cfg or {}).get("template_bindings") or {}
    direct_binding = template_bindings.get("teaching_load_summary") or {}
    nested_binding = (template_bindings.get("teaching_load") or {}).get("summary") or {}

    for binding in (nested_binding, direct_binding):
        raw_table_id = (binding or {}).get("raw_table_id")
        if not raw_table_id:
            continue
        raw_table = raw_tables.get(int(raw_table_id))
        if raw_table:
            return raw_table

    summary_tables = [
        raw_table
        for raw_table in (raw_tables or {}).values()
        if is_teaching_load_summary_raw_table(raw_table)
    ]
    if not summary_tables:
        return None

    summary_tables.sort(key=lambda item: int(item.get("table_index") or 0))
    return summary_tables[0]


def _find_summary_row_index(table, *patterns: str) -> Optional[int]:
    lowered_patterns = tuple(pattern.lower() for pattern in patterns if pattern)
    for row_index, row in enumerate(table.rows):
        first_cell_text = _normalize_text(row.cells[0].text if row.cells else "").lower()
        if not first_cell_text:
            continue
        if any(pattern in first_cell_text for pattern in lowered_patterns):
            return row_index
    return None


def _find_teaching_load_summary_row_indexes(table) -> Dict[str, int]:
    row_indexes = {
        "1": _find_summary_row_index(table, "1st term workload", "plan 1 sem"),
        "2": _find_summary_row_index(table, "2nd term workload", "plan 2 sem"),
        "annual": _find_summary_row_index(table, "academic year workload", "год.план"),
    }

    fallback_indexes = {"1": 2, "2": 3, "annual": 4}
    for key, fallback_index in fallback_indexes.items():
        if row_indexes.get(key) is None and len(table.rows) > fallback_index:
            row_indexes[key] = fallback_index

    return {key: value for key, value in row_indexes.items() if value is not None}


def _fill_teaching_load_summary_row(table, row_index: int, payload: Dict[str, Any]) -> None:
    for field_key, col_index in SUMMARY_TABLE_COLUMN_MAP.items():
        cell = _safe_get_cell(table, row_index, col_index)
        _set_cell_text(cell, _display_value((payload or {}).get(field_key)))


def _render_teaching_load_summary(
    doc: Document,
    raw_tables: Dict[int, Dict[str, Any]],
    settings_cfg: Dict[str, Any],
    context: Dict[str, Any],
) -> None:
    raw_table = _resolve_teaching_load_summary_raw_table(raw_tables, settings_cfg)
    if not raw_table:
        return

    table = _safe_get_table(doc, int(raw_table["table_index"]))
    if table is None:
        return

    row_indexes = _find_teaching_load_summary_row_indexes(table)
    if not row_indexes:
        return

    summary = build_teaching_load_summary((context.get("teaching_load") or {}), load_kind="staff")
    by_semester = summary.get("by_semester") or {}

    if row_indexes.get("1") is not None:
        _fill_teaching_load_summary_row(table, row_indexes["1"], by_semester.get("1") or {})
    if row_indexes.get("2") is not None:
        _fill_teaching_load_summary_row(table, row_indexes["2"], by_semester.get("2") or {})
    if row_indexes.get("annual") is not None:
        _fill_teaching_load_summary_row(table, row_indexes["annual"], summary.get("annual") or {})


def _render_teaching_load_for_kind(
    doc: Document,
    raw_tables: Dict[int, Dict[str, Any]],
    settings_cfg: Dict[str, Any],
    context: Dict[str, Any],
    load_kind: str,
):
    teaching_binding = (((settings_cfg or {}).get("template_bindings") or {}).get("teaching_load") or {}).get(load_kind) or {}
    raw_table_id = teaching_binding.get("raw_table_id")
    if not raw_table_id:
        return

    raw_table = raw_tables.get(int(raw_table_id))
    if not raw_table:
        return

    table = _safe_get_table(doc, raw_table["table_index"])
    if table is None:
        return

    blocks = _detect_table_blocks(table, raw_table)
    if not blocks:
        return

    col_map = _guess_column_map(raw_table)
    teaching_load = (context.get("teaching_load") or {})
    load_context = teaching_load.get(load_kind) or {}
    rows_by_scope = load_context.get("rows_by_scope") or {}
    mapped_rows = _map_scope_rows_to_blocks(
        rows_by_scope=rows_by_scope,
        blocks=blocks,
        primary_common_scope_key=teaching_load.get("primary_common_scope_key"),
    )

    for block in sorted(blocks, key=lambda item: item["total_row_index"], reverse=True):
        scope_key = block["scope_key"]
        scope_rows = mapped_rows.get(scope_key) or []
        _render_scope_block(
            table=table,
            block=block,
            rows_data=scope_rows,
            totals=_sum_scope_rows(scope_rows),
            col_map=col_map,
        )

    annual_total_row_index = _find_annual_total_row_index(
        table,
        max(block["total_row_index"] for block in blocks),
    )
    if annual_total_row_index is not None:
        _fill_total_row(
            table=table,
            row_index=annual_total_row_index,
            totals=load_context.get("annual_totals") or {},
            col_map=col_map,
        )


def _render_all_teaching_loads(
    doc: Document,
    raw_tables: Dict[int, Dict[str, Any]],
    settings_cfg: Dict[str, Any],
    context: Dict[str, Any],
):
    _render_teaching_load_for_kind(
        doc=doc,
        raw_tables=raw_tables,
        settings_cfg=settings_cfg,
        context=context,
        load_kind="staff",
    )
    _render_teaching_load_for_kind(
        doc=doc,
        raw_tables=raw_tables,
        settings_cfg=settings_cfg,
        context=context,
        load_kind="hourly",
    )
    _render_teaching_load_summary(
        doc=doc,
        raw_tables=raw_tables,
        settings_cfg=settings_cfg,
        context=context,
    )


def _build_output_path(teacher: Dict[str, Any], academic_year: str) -> str:
    safe_teacher_name = _safe_name(teacher["full_name"]) or f"teacher_{teacher['id']}"
    return str((Path(GENERATED_DIR) / f"IPP_{safe_teacher_name}_{academic_year}.docx").resolve())


def _load_generation_dependencies(cur, teacher_id: int, department_id: int, academic_year: str):
    teacher = _extract_teacher(cur, teacher_id)
    excel = _get_excel_by_year(cur, department_id, academic_year)
    raw_template = _get_raw_template_by_year(cur, department_id, academic_year)
    settings_cfg = _get_settings_for_excel(cur, excel["id"])
    excel_columns = _get_excel_columns(cur, excel["id"])
    excel_rows = _get_excel_rows(cur, excel["id"])
    raw_tables = _get_raw_tables(cur, raw_template["id"])

    return {
        "teacher": teacher,
        "excel": excel,
        "raw_template": raw_template,
        "settings_cfg": settings_cfg,
        "excel_columns": excel_columns,
        "excel_rows": excel_rows,
        "raw_tables": raw_tables,
        "academic_year": academic_year,
    }


def _build_generation_context(deps: Dict[str, Any]) -> Dict[str, Any]:
    return _build_excel_context(
        teacher=deps["teacher"],
        excel_columns=deps["excel_columns"],
        excel_rows=deps["excel_rows"],
        settings_cfg=deps["settings_cfg"],
        academic_year=deps["academic_year"],
    )


def _render_generated_doc(
    raw_template_path: str,
    raw_tables: Dict[int, Dict[str, Any]],
    settings_cfg: Dict[str, Any],
    context: Dict[str, Any],
) -> Document:
    doc = Document(raw_template_path)
    _render_all_teaching_loads(
        doc=doc,
        raw_tables=raw_tables,
        settings_cfg=settings_cfg,
        context=context,
    )
    return doc


def generate_docx_for_teacher(
    teacher_id: int,
    department_id: int,
    academic_year: str,
) -> str:
    conn = get_connection()
    try:
        with conn.cursor() as cur:
            deps = _load_generation_dependencies(
                cur=cur,
                teacher_id=teacher_id,
                department_id=department_id,
                academic_year=academic_year,
            )

        context = _build_generation_context(deps)

        doc = _render_generated_doc(
            raw_template_path=deps["raw_template"]["file_path"],
            raw_tables=deps["raw_tables"],
            settings_cfg=deps["settings_cfg"],
            context=context,
        )

        output_path = _build_output_path(deps["teacher"], academic_year)
        doc.save(output_path)

        apply_manual_fill_to_generated_docx(
            teacher_id=teacher_id,
            department_id=department_id,
            academic_year=academic_year,
            output_path=output_path,
        )

        return output_path
    finally:
        conn.close()
