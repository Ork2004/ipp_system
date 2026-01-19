import re
from dataclasses import dataclass
from typing import Optional, List, Dict, Set, Tuple, Any

from docx import Document

TOKEN_PATTERN = re.compile(r"({%\s*for\s+.*?%}|{%\s*endfor\s*%}|{{\s*.*?\s*}})")
LOOP_START_PATTERN = re.compile(r"{%\s*for\s+(\w+)\s+in\s+(\w+)\s*%}")
LOOP_END_PATTERN = re.compile(r"{%\s*endfor\s*%}")
TEXT_PATTERN = re.compile(r"{{\s*(.*?)\s*}}")


@dataclass
class LoopContext:
    var: str
    array: str


def _make_path_paragraph(index: int) -> str:
    return f"paragraph[{index}]"


def _make_path_table(t: int, r: int, c: int) -> str:
    return f"table[{t}]/row[{r}]/cell[{c}]"


def analyze_docx(file_path: str) -> tuple[list[dict], list[dict]]:
    doc = Document(file_path)

    placeholder_schema: List[Dict[str, Any]] = []
    schema_seen: Set[Tuple] = set()

    placeholders: List[Dict[str, Any]] = []
    placeholders_seen: Set[Tuple] = set()

    loop_stack: List[LoopContext] = []

    def register_loop(array_name: str, var_name: str, path: str):
        schema_key = (array_name, "loop", path, var_name)
        if schema_key not in schema_seen:
            schema_seen.add(schema_key)
            placeholder_schema.append({
                "name": array_name,
                "type": "loop",
                "loop_var": var_name,
                "path": path
            })

        ph_key = (array_name, "loop")
        if ph_key not in placeholders_seen:
            placeholders_seen.add(ph_key)
            placeholders.append({
                "placeholder_name": array_name,
                "placeholder_type": "loop",
                "extra_meta": {"loop_var": var_name}
            })

        loop_stack.append(LoopContext(var=var_name, array=array_name))

    def register_text(placeholder_name: str, path: str):
        loop_array: Optional[str] = None
        loop_var: Optional[str] = None
        if loop_stack:
            loop_array = loop_stack[-1].array
            loop_var = loop_stack[-1].var

        schema_key = (placeholder_name, "text", path, loop_array)
        if schema_key not in schema_seen:
            schema_seen.add(schema_key)
            item: Dict[str, Any] = {"name": placeholder_name, "type": "text", "path": path}
            if loop_array is not None:
                item["loop"] = loop_array
                item["loop_var"] = loop_var
            placeholder_schema.append(item)

        ph_key = (placeholder_name, "text", loop_array)
        if ph_key not in placeholders_seen:
            placeholders_seen.add(ph_key)
            extra_meta = None
            if loop_array is not None:
                extra_meta = {"loop": loop_array, "loop_var": loop_var}

            placeholders.append({
                "placeholder_name": placeholder_name,
                "placeholder_type": "text",
                "extra_meta": extra_meta
            })

    def process_text(text: str, path: str):
        for token_match in TOKEN_PATTERN.finditer(text):
            token = token_match.group(0)

            m_for = LOOP_START_PATTERN.match(token)
            if m_for:
                var_name, array_name = m_for.group(1), m_for.group(2)
                register_loop(array_name=array_name, var_name=var_name, path=path)
                continue

            if LOOP_END_PATTERN.match(token):
                if loop_stack:
                    loop_stack.pop()
                continue

            m_text = TEXT_PATTERN.match(token)
            if m_text:
                placeholder_name = m_text.group(1).strip()
                if placeholder_name:
                    register_text(placeholder_name=placeholder_name, path=path)

    for p_index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text:
            process_text(paragraph.text, _make_path_paragraph(p_index))

    for t_index, table in enumerate(doc.tables):
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        process_text(paragraph.text, _make_path_table(t_index, r_index, c_index))

    return placeholder_schema, placeholders
