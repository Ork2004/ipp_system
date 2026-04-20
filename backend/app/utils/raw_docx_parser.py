import hashlib
import re
from typing import Any, Dict, List, Optional

from docx import Document


SECTION_HEADING_RE = re.compile(
    r"^\s*((?:\d+(\.\d+)*)|(?:[IVXLC]+))[\).\s-]+.+",
    re.IGNORECASE
)

TOTAL_ROW_RE = re.compile(
    r"(итого|барлығы|total|всего)",
    re.IGNORECASE
)

NUMBER_ONLY_RE = re.compile(r"^\s*\d+\s*$")


def _normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _normalize_text_lower(value: Any) -> str:
    return _normalize_text(value).lower()


def _is_empty_text(value: Any) -> bool:
    return _normalize_text(value) == ""


def _is_section_heading(text: str) -> bool:
    text = _normalize_text(text)
    if not text:
        return False
    return bool(SECTION_HEADING_RE.match(text))


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _cell_signature(text: str) -> str:
    text = _normalize_text(text).lower()
    if not text:
        return "__empty__"
    text = re.sub(r"\d+", "#", text)
    return text


def _build_header_signature(matrix: List[List[dict]], take_rows: int = 2) -> str:
    parts: List[str] = []
    max_rows = min(len(matrix), take_rows)

    for r in range(max_rows):
        row_parts = []
        for cell in matrix[r]:
            row_parts.append(_cell_signature(cell.get("text")))
        parts.append("|".join(row_parts))

    return " || ".join(parts)


def _row_has_total_label(row: List[dict]) -> bool:
    for cell in row:
        text = _normalize_text(cell.get("text"))
        if text and TOTAL_ROW_RE.search(text):
            return True
    return False


def _is_probable_loop_template_row(row: List[dict]) -> bool:
    if not row:
        return False

    texts = [_normalize_text(c.get("text")) for c in row]
    non_empty_indices = [i for i, t in enumerate(texts) if t != ""]
    empty_count = sum(1 for t in texts if t == "")

    if len(row) < 2:
        return False

    if empty_count >= max(2, len(row) - 1):
        return True

    if texts and NUMBER_ONLY_RE.match(texts[0] or ""):
        rest = texts[1:]
        if rest and sum(1 for t in rest if t == "") >= max(1, len(rest) - 1):
            return True

    if len(non_empty_indices) == 1 and non_empty_indices[0] == 0:
        return True

    return False


def _guess_table_type(matrix: List[List[dict]]) -> str:
    if not matrix:
        return "static"

    body_rows = matrix[1:] if len(matrix) > 1 else matrix

    probable_loop_rows = 0
    total_rows = 0

    for row in body_rows:
        if _row_has_total_label(row):
            continue
        total_rows += 1
        if _is_probable_loop_template_row(row):
            probable_loop_rows += 1

    if total_rows > 0 and probable_loop_rows >= 1:
        return "loop"

    return "static"


def _extract_paragraphs_with_order(doc: Document) -> List[dict]:
    out: List[dict] = []
    for i, p in enumerate(doc.paragraphs):
        text = _normalize_text(p.text)
        out.append({
            "type": "paragraph",
            "index": i,
            "text": text,
        })
    return out


def _find_table_section_title(doc: Document, table_index: int) -> str:
    section_title = ""
    paragraphs = _extract_paragraphs_with_order(doc)

    for p in paragraphs:
        txt = p["text"]
        if _is_section_heading(txt):
            section_title = txt

    if section_title:
        return section_title

    return f"Таблица {table_index + 1}"


def _guess_column_hint_text(matrix: List[List[dict]], col_index: int) -> str:
    if not matrix:
        return f"Колонка {col_index + 1}"

    for row in matrix[:2]:
        if col_index >= len(row):
            continue
        txt = _normalize_text(row[col_index].get("text"))
        if txt:
            return txt

    return f"Колонка {col_index + 1}"


def _build_semantic_key(
    *,
    section_title: str,
    table_type: str,
    header_signature: str,
    row_index: int,
    col_index: int,
    column_hint_text: str,
) -> str:
    base = " | ".join([
        _normalize_text_lower(section_title),
        _normalize_text_lower(table_type),
        _normalize_text_lower(header_signature),
        str(row_index),
        str(col_index),
        _normalize_text_lower(column_hint_text),
    ])
    return hashlib.sha1(base.encode("utf-8")).hexdigest()


def _build_table_fingerprint(
    *,
    section_title: str,
    table_type: str,
    header_signature: str,
) -> str:
    base = " | ".join([
        _normalize_text_lower(section_title),
        _normalize_text_lower(table_type),
        _normalize_text_lower(header_signature),
    ])
    return hashlib.sha1(base.encode("utf-8")).hexdigest()


def _build_structure_meta(
    *,
    section_title: str,
    table_type: str,
    header_signature: str,
    column_hints: List[str],
    row_count: int,
    col_count: int,
    has_total_row: bool,
) -> Dict[str, Any]:
    return {
        "section_title_norm": _normalize_text_lower(section_title),
        "table_type_norm": _normalize_text_lower(table_type),
        "header_signature_norm": _normalize_text_lower(header_signature),
        "column_hints_norm": [_normalize_text_lower(x) for x in (column_hints or [])],
        "row_count": row_count,
        "col_count": col_count,
        "has_total_row": bool(has_total_row),
    }


def _extract_table_matrix(
    table,
    table_index: int,
    *,
    section_title: str,
    table_type: str,
    header_signature: str,
) -> List[List[dict]]:
    matrix: List[List[dict]] = []

    for r_idx, row in enumerate(table.rows):
        row_cells: List[dict] = []

        for c_idx, cell in enumerate(row.cells):
            text = _normalize_text(cell.text)
            column_hint_text = ""

            row_cells.append({
                "row_index": r_idx,
                "col_index": c_idx,
                "text": text,
                "is_empty": _is_empty_text(text),
                "editable": _is_empty_text(text),
                "cell_key": f"t{table_index}_r{r_idx}_c{c_idx}",
                "column_hint_text": column_hint_text,
                "row_signature": f"r{r_idx}",
                "semantic_key": "",
                "cell_kind": "text",
            })

        matrix.append(row_cells)

    for r_idx, row in enumerate(matrix):
        for c_idx, cell in enumerate(row):
            column_hint_text = _guess_column_hint_text(matrix, c_idx)
            cell["column_hint_text"] = column_hint_text
            cell["semantic_key"] = _build_semantic_key(
                section_title=section_title,
                table_type=table_type,
                header_signature=header_signature,
                row_index=r_idx,
                col_index=c_idx,
                column_hint_text=column_hint_text,
            )

    return matrix


def _collect_editable_cells(matrix: List[List[dict]]) -> List[dict]:
    editable: List[dict] = []
    for row in matrix:
        for cell in row:
            if cell.get("editable"):
                editable.append({
                    "row_index": cell["row_index"],
                    "col_index": cell["col_index"],
                    "cell_key": cell["cell_key"],
                    "semantic_key": cell.get("semantic_key"),
                    "row_signature": cell.get("row_signature"),
                    "column_hint_text": cell.get("column_hint_text"),
                })
    return editable


def _collect_prefilled_cells(matrix: List[List[dict]]) -> List[dict]:
    prefilled: List[dict] = []
    for row in matrix:
        for cell in row:
            if not cell.get("editable"):
                prefilled.append({
                    "row_index": cell["row_index"],
                    "col_index": cell["col_index"],
                    "text": cell["text"],
                    "cell_key": cell["cell_key"],
                    "semantic_key": cell.get("semantic_key"),
                    "row_signature": cell.get("row_signature"),
                    "column_hint_text": cell.get("column_hint_text"),
                })
    return prefilled


def _extract_loop_template_row_index(matrix: List[List[dict]]) -> Optional[int]:
    for row in matrix:
        if _row_has_total_label(row):
            continue
        if _is_probable_loop_template_row(row):
            return _safe_int(row[0].get("row_index"))
    return None


def _extract_column_hints(matrix: List[List[dict]]) -> List[str]:
    if not matrix:
        return []

    first_row = matrix[0]
    hints: List[str] = []

    for idx, cell in enumerate(first_row):
        txt = _normalize_text(cell.get("text"))
        if txt:
            hints.append(txt)
        else:
            hints.append(f"Колонка {idx + 1}")

    return hints


def scan_raw_docx(file_path: str) -> Dict[str, Any]:
    doc = Document(file_path)
    tables_out: List[Dict[str, Any]] = []

    for t_idx, table in enumerate(doc.tables):
        section_title = _find_table_section_title(doc, t_idx)

        temp_matrix = []
        for r_idx, row in enumerate(table.rows):
            row_cells: List[dict] = []
            for c_idx, cell in enumerate(row.cells):
                text = _normalize_text(cell.text)
                row_cells.append({
                    "row_index": r_idx,
                    "col_index": c_idx,
                    "text": text,
                    "is_empty": _is_empty_text(text),
                    "editable": _is_empty_text(text),
                    "cell_key": f"t{t_idx}_r{r_idx}_c{c_idx}",
                })
            temp_matrix.append(row_cells)

        row_count = len(temp_matrix)
        col_count = max((len(r) for r in temp_matrix), default=0)

        table_type = _guess_table_type(temp_matrix)
        header_signature = _build_header_signature(temp_matrix)
        column_hints = _extract_column_hints(temp_matrix)
        has_total_row = any(_row_has_total_label(row) for row in temp_matrix)
        loop_template_row_index = _extract_loop_template_row_index(temp_matrix)

        table_fingerprint = _build_table_fingerprint(
            section_title=section_title,
            table_type=table_type,
            header_signature=header_signature,
        )

        structure_meta = _build_structure_meta(
            section_title=section_title,
            table_type=table_type,
            header_signature=header_signature,
            column_hints=column_hints,
            row_count=row_count,
            col_count=col_count,
            has_total_row=has_total_row,
        )

        matrix = _extract_table_matrix(
            table,
            t_idx,
            section_title=section_title,
            table_type=table_type,
            header_signature=header_signature,
        )

        editable_cells = _collect_editable_cells(matrix)
        prefilled_cells = _collect_prefilled_cells(matrix)

        tables_out.append({
            "table_index": t_idx,
            "section_title": section_title,
            "table_type": table_type,
            "row_count": row_count,
            "col_count": col_count,
            "header_signature": header_signature,
            "has_total_row": has_total_row,
            "loop_template_row_index": loop_template_row_index,
            "column_hints": column_hints,
            "editable_cells_count": len(editable_cells),
            "prefilled_cells_count": len(prefilled_cells),
            "table_fingerprint": table_fingerprint,
            "structure_meta": structure_meta,
            "editable_cells": editable_cells,
            "prefilled_cells": prefilled_cells,
            "matrix": matrix,
        })

    return {
        "file_path": file_path,
        "tables_count": len(tables_out),
        "tables": tables_out,
    }