import re
from typing import Any, Dict, List, Optional

from docx import Document
from psycopg2.extras import Json

from backend.app.database import get_connection
from backend.app.utils.teaching_load import (
    build_effective_generation_settings,
    extract_excel_bound_raw_table_ids_with_raw_tables,
)


TOTAL_ROW_VARIANTS = ("итого", "total", "всего", "барлығы")


def _normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _normalize_text_lower(value: Any) -> str:
    return _normalize_text(value).lower()


def _safe_cell_text(table, row_index: int, col_index: int) -> str:
    try:
        return _normalize_text(table.rows[row_index].cells[col_index].text)
    except Exception:
        return ""


def _is_total_row(texts: List[str]) -> bool:
    return any(
        any(marker in _normalize_text_lower(text) for marker in TOTAL_ROW_VARIANTS)
        for text in texts
        if _normalize_text(text)
    )


def _is_placeholder_loop_row(texts: List[str]) -> bool:
    non_empty = [text for text in texts if _normalize_text(text)]
    if not non_empty:
        return True
    if len(non_empty) == 1 and re.fullmatch(r"\d+", non_empty[0]):
        return True
    return False


def _is_header_like_row(texts: List[str], column_hints: List[str]) -> bool:
    normalized_texts = [_normalize_text_lower(text) for text in texts if _normalize_text(text)]
    normalized_hints = [_normalize_text_lower(text) for text in (column_hints or []) if _normalize_text(text)]

    if not normalized_texts or not normalized_hints:
        return False

    matches = 0
    for text in normalized_texts:
        if any(text == hint or text in hint or hint in text for hint in normalized_hints):
            matches += 1

    return matches >= max(2, len(normalized_texts) - 1)


def _find_total_row_index(table, start_from: int) -> Optional[int]:
    for row_index in range(max(start_from, 0), len(table.rows)):
        texts = [_normalize_text(cell.text) for cell in table.rows[row_index].cells]
        if _is_total_row(texts):
            return row_index
    return None


def _looks_like_placeholder_cell(text: str) -> bool:
    normalized = _normalize_text_lower(text)
    if not normalized:
        return False
    if "___" in text or "20___" in text or "20__" in text:
        return True
    return any(token in normalized for token in ("ф.и.о", "full name", "approved by"))


def _load_raw_template(cur, department_id: int, academic_year: str) -> Optional[Dict[str, Any]]:
    cur.execute(
        """
        SELECT id
        FROM raw_docx_templates
        WHERE department_id = %s
          AND academic_year = %s
        LIMIT 1;
        """,
        (department_id, academic_year),
    )
    row = cur.fetchone()
    if not row:
        return None
    return {"id": int(row[0])}


def _load_generation_settings_config(cur, department_id: int, academic_year: str) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT gs.config
        FROM generation_settings gs
        JOIN excel_templates et ON et.id = gs.excel_template_id
        WHERE et.department_id = %s
          AND et.academic_year = %s
        LIMIT 1;
        """,
        (department_id, academic_year),
    )
    row = cur.fetchone()
    if not row:
        return {}
    return row[0] or {}


def _load_raw_tables(cur, raw_template_id: int) -> List[Dict[str, Any]]:
    cur.execute(
        """
            SELECT
                id,
                table_index,
            section_title,
            table_type,
            row_count,
            col_count,
            header_signature,
            has_total_row,
            loop_template_row_index,
            column_hints,
            table_fingerprint
        FROM raw_docx_tables
        WHERE template_id = %s
        ORDER BY table_index;
        """,
        (raw_template_id,),
    )
    table_rows = cur.fetchall() or []
    out: List[Dict[str, Any]] = []

    for row in table_rows:
        raw_table = {
            "id": int(row[0]),
            "table_index": int(row[1] or 0),
            "section_title": row[2] or "",
            "table_type": row[3] or "",
            "row_count": int(row[4] or 0),
            "col_count": int(row[5] or 0),
            "header_signature": row[6] or "",
            "has_total_row": bool(row[7]),
            "loop_template_row_index": row[8] if row[8] is None else int(row[8]),
            "column_hints": row[9] or [],
            "table_fingerprint": row[10] or "",
            "editable_cells": [],
        }

        cur.execute(
            """
            SELECT
                id,
                row_index,
                col_index,
                cell_key,
                original_text,
                semantic_key,
                row_signature,
                column_hint_text,
                is_editable
            FROM raw_docx_cells
            WHERE table_id = %s
            ORDER BY row_index, col_index;
            """,
            (raw_table["id"],),
        )
        for cell_row in cur.fetchall() or []:
            original_text = cell_row[4] or ""
            is_fillable = bool(cell_row[8]) or (
                raw_table["table_type"] == "static" and _looks_like_placeholder_cell(original_text)
            )
            if not is_fillable:
                continue
            raw_table["editable_cells"].append(
                {
                    "raw_cell_id": int(cell_row[0]),
                    "row_index": int(cell_row[1]),
                    "col_index": int(cell_row[2]),
                    "cell_key": cell_row[3],
                    "semantic_key": cell_row[5],
                    "row_signature": cell_row[6],
                    "column_hint_text": cell_row[7],
                }
            )

        out.append(raw_table)

    return out


def _delete_current_snapshot(cur, teacher_id: int, academic_year: str, raw_table_id: int) -> None:
    cur.execute(
        """
        DELETE FROM teacher_manual_table_snapshots
        WHERE teacher_id = %s
          AND academic_year = %s
          AND raw_table_id = %s;
        """,
        (teacher_id, academic_year, raw_table_id),
    )


def _create_snapshot(
    cur,
    *,
    teacher_id: int,
    department_id: int,
    academic_year: str,
    raw_template_id: int,
    raw_table: Dict[str, Any],
) -> int:
    cur.execute(
        """
        INSERT INTO teacher_manual_table_snapshots(
            teacher_id,
            academic_year,
            raw_template_id,
            raw_table_id,
            department_id,
            section_title,
            table_type,
            header_signature,
            column_hints,
            table_fingerprint,
            source_mode,
            prefilled_from_snapshot_id,
            created_at,
            updated_at
        )
        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'manual',NULL,now(),now())
        RETURNING id;
        """,
        (
            teacher_id,
            academic_year,
            raw_template_id,
            raw_table["id"],
            department_id,
            raw_table.get("section_title"),
            raw_table.get("table_type"),
            raw_table.get("header_signature"),
            Json(raw_table.get("column_hints") or []),
            raw_table.get("table_fingerprint"),
        ),
    )
    return int(cur.fetchone()[0])


def _extract_static_values(raw_table: Dict[str, Any], filled_table) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for cell in raw_table.get("editable_cells") or []:
        value = _safe_cell_text(filled_table, int(cell["row_index"]), int(cell["col_index"]))
        if not value:
            continue
        out.append({**cell, "value": value})
    return out


def _extract_loop_rows(raw_table: Dict[str, Any], filled_table) -> List[Dict[str, Any]]:
    start_row = int(raw_table.get("loop_template_row_index") or 0)
    total_row_index = _find_total_row_index(filled_table, start_row)
    end_row = total_row_index if total_row_index is not None else len(filled_table.rows)
    col_count = int(raw_table.get("col_count") or 0)
    column_hints = raw_table.get("column_hints") or []

    out: List[Dict[str, Any]] = []
    for row_index in range(start_row, end_row):
        texts = [_safe_cell_text(filled_table, row_index, col_index) for col_index in range(col_count)]
        if _is_total_row(texts):
            continue
        if _is_header_like_row(texts, column_hints):
            continue
        if _is_placeholder_loop_row(texts):
            continue

        values = []
        for col_index, text in enumerate(texts):
            if not text:
                continue
            values.append(
                {
                    "col_index": col_index,
                    "column_hint_text": column_hints[col_index] if col_index < len(column_hints) else f"Колонка {col_index + 1}",
                    "value": text,
                }
            )
        if not values:
            continue

        out.append(
            {
                "row_order": len(out) + 1,
                "values": values,
            }
        )

    return out


def import_manual_docx(
    *,
    file_path: str,
    teacher_id: int,
    department_id: int,
    academic_year: str,
) -> Dict[str, Any]:
    doc = Document(file_path)

    conn = get_connection()
    try:
        with conn:
            with conn.cursor() as cur:
                raw_template = _load_raw_template(cur, department_id, academic_year)
                if not raw_template:
                    raise Exception("Raw шаблон для этого года не найден")

                raw_tables = _load_raw_tables(cur, raw_template["id"])
                settings_cfg = _load_generation_settings_config(cur, department_id, academic_year)
                effective_settings = build_effective_generation_settings(settings_cfg, raw_tables)
                excel_bound_ids = extract_excel_bound_raw_table_ids_with_raw_tables(
                    effective_settings,
                    raw_tables,
                )

                imported_static_tables = 0
                imported_loop_tables = 0
                imported_static_values = 0
                imported_loop_rows = 0

                for raw_table in raw_tables:
                    if int(raw_table["id"]) in excel_bound_ids:
                        continue
                    if int(raw_table["table_index"]) >= len(doc.tables):
                        continue

                    filled_table = doc.tables[int(raw_table["table_index"])]
                    _delete_current_snapshot(cur, teacher_id, academic_year, int(raw_table["id"]))
                    snapshot_id = _create_snapshot(
                        cur,
                        teacher_id=teacher_id,
                        department_id=department_id,
                        academic_year=academic_year,
                        raw_template_id=int(raw_template["id"]),
                        raw_table=raw_table,
                    )

                    if raw_table.get("table_type") == "static":
                        static_values = _extract_static_values(raw_table, filled_table)
                        for item in static_values:
                            cur.execute(
                                """
                                INSERT INTO teacher_manual_static_cell_values(
                                    snapshot_id,
                                    raw_cell_id,
                                    row_index,
                                    col_index,
                                    cell_key,
                                    semantic_key,
                                    row_signature,
                                    column_hint_text,
                                    value_text,
                                    created_at,
                                    updated_at
                                )
                                VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,now(),now());
                                """,
                                (
                                    snapshot_id,
                                    item["raw_cell_id"],
                                    item["row_index"],
                                    item["col_index"],
                                    item["cell_key"],
                                    item["semantic_key"],
                                    item["row_signature"],
                                    item["column_hint_text"],
                                    item["value"],
                                ),
                            )
                        imported_static_tables += 1
                        imported_static_values += len(static_values)
                        continue

                    loop_rows = _extract_loop_rows(raw_table, filled_table)
                    for row in loop_rows:
                        cur.execute(
                            """
                            INSERT INTO teacher_manual_loop_rows(
                                snapshot_id,
                                row_order,
                                created_at,
                                updated_at
                            )
                            VALUES (%s,%s,now(),now())
                            RETURNING id;
                            """,
                            (snapshot_id, row["row_order"]),
                        )
                        loop_row_id = int(cur.fetchone()[0])

                        for cell in row["values"]:
                            cur.execute(
                                """
                                INSERT INTO teacher_manual_loop_cell_values(
                                    loop_row_id,
                                    col_index,
                                    column_hint_text,
                                    semantic_key,
                                    value_text,
                                    created_at,
                                    updated_at
                                )
                                VALUES (%s,%s,%s,%s,%s,now(),now());
                                """,
                                (
                                    loop_row_id,
                                    cell["col_index"],
                                    cell["column_hint_text"],
                                    None,
                                    cell["value"],
                                ),
                            )

                    imported_loop_tables += 1
                    imported_loop_rows += len(loop_rows)

        return {
            "status": "ok",
            "imported_static_tables": imported_static_tables,
            "imported_loop_tables": imported_loop_tables,
            "imported_static_values": imported_static_values,
            "imported_loop_rows": imported_loop_rows,
        }
    finally:
        conn.close()
