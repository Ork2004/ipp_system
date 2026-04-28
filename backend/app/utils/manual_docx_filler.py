from copy import deepcopy
from typing import Any, Dict, List, Optional

from docx import Document

from backend.app.database import get_connection
from backend.app.utils.teaching_load import (
    extract_excel_bound_raw_table_ids,
    get_teaching_load_summary_binding,
    is_manual_source_binding,
    is_teaching_load_summary_raw_table,
)


def _normalize_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _clear_cell(cell):
    cell.text = ""


def _set_cell_text(cell, value: Any):
    _clear_cell(cell)
    cell.text = "" if value is None else str(value)


def _safe_get_table(doc: Document, table_index: int):
    if table_index < 0 or table_index >= len(doc.tables):
        return None
    return doc.tables[table_index]


def _safe_get_cell(table, row_index: int, col_index: int):
    try:
        row = table.rows[row_index]
        cell = row.cells[col_index]
        return cell
    except Exception:
        return None


def _insert_row_after(table, row_index: int):
    tr = table.rows[row_index]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)
    return row_index + 1


def _insert_row_before(table, row_index: int):
    tr = table.rows[row_index]._tr
    new_tr = deepcopy(tr)
    tr.addprevious(new_tr)
    return row_index


def _remove_row(table, row_index: int):
    tr = table.rows[row_index]._tr
    tbl = tr.getparent()
    tbl.remove(tr)


def _find_total_row_index(table, start_from: int = 0) -> Optional[int]:
    keywords = ("итого", "total", "всего", "барлығы")
    for r_idx in range(start_from, len(table.rows)):
        row = table.rows[r_idx]
        for cell in row.cells:
            txt = _normalize_text(cell.text).lower()
            if txt and any(k in txt for k in keywords):
                return r_idx
    return None


def _load_raw_template_id(cur, department_id: int, academic_year: str) -> Optional[int]:
    cur.execute("""
        SELECT id
        FROM raw_docx_templates
        WHERE department_id=%s AND academic_year=%s
        LIMIT 1;
    """, (department_id, academic_year))
    row = cur.fetchone()
    return row[0] if row else None


def _load_snapshot_map(cur, teacher_id: int, academic_year: str) -> Dict[int, Dict[str, Any]]:
    cur.execute("""
        SELECT
            id,
            raw_table_id,
            table_type,
            source_mode,
            prefilled_from_snapshot_id
        FROM teacher_manual_table_snapshots
        WHERE teacher_id=%s
          AND academic_year=%s
        ORDER BY raw_table_id, updated_at DESC, id DESC;
    """, (teacher_id, academic_year))
    rows = cur.fetchall() or []

    out: Dict[int, Dict[str, Any]] = {}
    for r in rows:
        raw_table_id = int(r[1])
        if raw_table_id not in out:
            out[raw_table_id] = {
                "snapshot_id": int(r[0]),
                "raw_table_id": raw_table_id,
                "table_type": r[2],
                "source_mode": r[3],
                "prefilled_from_snapshot_id": r[4],
            }
    return out


def _load_generation_settings_config(cur, department_id: int, academic_year: str) -> Dict[str, Any]:
    cur.execute(
        """
        SELECT gs.config
        FROM generation_settings gs
        JOIN excel_templates et ON et.id = gs.excel_template_id
        WHERE et.department_id = %s
          AND et.academic_year = %s
        LIMIT 1;
        """,
        (department_id, academic_year),
    )
    row = cur.fetchone()
    if not row:
        return {}
    return row[0] or {}


def _load_auto_excel_bound_raw_table_ids(
    cur,
    raw_template_id: int,
    *,
    include_summary_tables: bool,
) -> set[int]:
    if not include_summary_tables:
        return set()

    cur.execute(
        """
        SELECT
            id,
            table_index,
            table_type,
            row_count,
            col_count,
            column_hints
        FROM raw_docx_tables
        WHERE template_id = %s;
        """,
        (raw_template_id,),
    )
    out: set[int] = set()

    for row in cur.fetchall() or []:
        raw_table = {
            "id": int(row[0]),
            "table_index": int(row[1] or 0),
            "table_type": row[2] or "",
            "row_count": int(row[3] or 0),
            "col_count": int(row[4] or 0),
            "column_hints": row[5] or [],
        }
        if is_teaching_load_summary_raw_table(raw_table):
            out.add(int(row[0]))

    return out


def _load_manual_static_tables(
    cur,
    raw_template_id: int,
    teacher_id: int,
    academic_year: str,
    excluded_raw_table_ids: set[int],
) -> List[Dict[str, Any]]:
    cur.execute("""
        SELECT
            t.id,
            t.table_index,
            t.section_title,
            t.table_type,
            t.row_count,
            t.col_count
        FROM raw_docx_tables t
        WHERE t.template_id=%s
          AND t.table_type='static'
        ORDER BY t.table_index;
    """, (raw_template_id,))
    tables = cur.fetchall() or []

    snapshot_map = _load_snapshot_map(cur, teacher_id, academic_year)
    out = []

    for t in tables:
        raw_table_id = int(t[0])
        if raw_table_id in excluded_raw_table_ids:
            continue
        snapshot = snapshot_map.get(raw_table_id)

        editable_values = []
        if snapshot and str(snapshot["table_type"]) == "static":
            cur.execute("""
                SELECT
                    raw_cell_id,
                    row_index,
                    col_index,
                    value_text
                FROM teacher_manual_static_cell_values
                WHERE snapshot_id=%s
                ORDER BY row_index, col_index, id;
            """, (snapshot["snapshot_id"],))
            vals = cur.fetchall() or []

            editable_values = [
                {
                    "raw_cell_id": v[0],
                    "row_index": v[1],
                    "col_index": v[2],
                    "value": v[3] if v[3] is not None else "",
                }
                for v in vals
            ]

        out.append({
            "raw_table_id": raw_table_id,
            "table_index": t[1],
            "section_title": t[2],
            "table_type": t[3],
            "row_count": t[4],
            "col_count": t[5],
            "editable_values": editable_values,
        })

    return out


def _load_manual_loop_tables(
    cur,
    raw_template_id: int,
    teacher_id: int,
    academic_year: str,
    excluded_raw_table_ids: set[int],
) -> List[Dict[str, Any]]:
    cur.execute("""
        SELECT
            t.id,
            t.table_index,
            t.section_title,
            t.table_type,
            t.row_count,
            t.col_count,
            t.loop_template_row_index,
            t.has_total_row,
            t.column_hints
        FROM raw_docx_tables t
        WHERE t.template_id=%s
          AND t.table_type='loop'
        ORDER BY t.table_index;
    """, (raw_template_id,))
    tables = cur.fetchall() or []

    snapshot_map = _load_snapshot_map(cur, teacher_id, academic_year)
    out = []

    for t in tables:
        raw_table_id = int(t[0])
        if raw_table_id in excluded_raw_table_ids:
            continue
        snapshot = snapshot_map.get(raw_table_id)

        rows_out = []
        if snapshot and str(snapshot["table_type"]) == "loop":
            cur.execute("""
                SELECT id, row_order
                FROM teacher_manual_loop_rows
                WHERE snapshot_id=%s
                ORDER BY row_order, id;
            """, (snapshot["snapshot_id"],))
            loop_rows = cur.fetchall() or []

            for lr in loop_rows:
                loop_row_id = lr[0]

                cur.execute("""
                    SELECT col_index, value_text
                    FROM teacher_manual_loop_cell_values
                    WHERE loop_row_id=%s
                    ORDER BY col_index, id;
                """, (loop_row_id,))
                vals = cur.fetchall() or []

                rows_out.append({
                    "loop_row_id": loop_row_id,
                    "row_order": lr[1],
                    "values": [
                        {
                            "col_index": v[0],
                            "value": v[1] if v[1] is not None else "",
                        }
                        for v in vals
                    ],
                })

        out.append({
            "raw_table_id": raw_table_id,
            "table_index": t[1],
            "section_title": t[2],
            "table_type": t[3],
            "row_count": t[4],
            "col_count": t[5],
            "loop_template_row_index": t[6],
            "has_total_row": t[7],
            "column_hints": t[8] or [],
            "loop_rows": rows_out,
        })

    return out


def _apply_static_table_values(doc: Document, table_item: Dict[str, Any]):
    table = _safe_get_table(doc, int(table_item["table_index"]))
    if table is None:
        return

    for item in table_item.get("editable_values", []):
        row_index = int(item["row_index"])
        col_index = int(item["col_index"])
        value = item.get("value", "")

        cell = _safe_get_cell(table, row_index, col_index)
        if cell is None:
            continue

        _set_cell_text(cell, value)


def _fill_row_cells(row, col_count: int, values_by_col: Dict[int, Any]):
    for col_index in range(col_count):
        try:
            cell = row.cells[col_index]
        except Exception:
            continue
        value = values_by_col.get(col_index, "")
        _set_cell_text(cell, value)


def _apply_loop_table_values(doc: Document, table_item: Dict[str, Any]):
    table = _safe_get_table(doc, int(table_item["table_index"]))
    if table is None:
        return

    template_row_index = table_item.get("loop_template_row_index")
    if template_row_index is None:
        return

    template_row_index = int(template_row_index)

    if template_row_index < 0 or template_row_index >= len(table.rows):
        return

    loop_rows = table_item.get("loop_rows", []) or []
    col_count = int(table_item.get("col_count") or 0)

    if not loop_rows:
        return

    total_row_index = _find_total_row_index(table, start_from=template_row_index + 1)

    inserted_row_indexes = []

    for loop_row in loop_rows:
        values_by_col = {
            int(v["col_index"]): v.get("value", "")
            for v in loop_row.get("values", [])
        }

        if total_row_index is not None:
            new_index = _insert_row_before(table, total_row_index)
            inserted_row_indexes.append(new_index)
            total_row_index += 1
        else:
            anchor_index = inserted_row_indexes[-1] if inserted_row_indexes else template_row_index
            new_index = _insert_row_after(table, anchor_index)
            inserted_row_indexes.append(new_index)

        try:
            new_row = table.rows[new_index]
        except Exception:
            continue

        _fill_row_cells(new_row, col_count, values_by_col)

    try:
        _remove_row(table, template_row_index)
    except Exception:
        pass


def apply_manual_fill_to_generated_docx(
    *,
    teacher_id: int,
    department_id: int,
    academic_year: str,
    output_path: str,
) -> None:
    conn = get_connection()
    try:
        with conn.cursor() as cur:
            raw_template_id = _load_raw_template_id(cur, department_id, academic_year)
            if not raw_template_id:
                return

            settings_cfg = _load_generation_settings_config(cur, department_id, academic_year)
            excel_bound_raw_table_ids = extract_excel_bound_raw_table_ids(settings_cfg)
            summary_binding = get_teaching_load_summary_binding(settings_cfg)
            include_auto_summary_tables = not (
                summary_binding and is_manual_source_binding(summary_binding)
            ) and not summary_binding.get("raw_table_id")

            excel_bound_raw_table_ids.update(
                _load_auto_excel_bound_raw_table_ids(
                    cur,
                    raw_template_id,
                    include_summary_tables=include_auto_summary_tables,
                )
            )

            static_tables = _load_manual_static_tables(
                cur=cur,
                raw_template_id=raw_template_id,
                teacher_id=teacher_id,
                academic_year=academic_year,
                excluded_raw_table_ids=excel_bound_raw_table_ids,
            )
            loop_tables = _load_manual_loop_tables(
                cur=cur,
                raw_template_id=raw_template_id,
                teacher_id=teacher_id,
                academic_year=academic_year,
                excluded_raw_table_ids=excel_bound_raw_table_ids,
            )

        doc = Document(output_path)

        for table_item in static_tables:
            _apply_static_table_values(doc, table_item)

        for table_item in loop_tables:
            _apply_loop_table_values(doc, table_item)

        doc.save(output_path)

    finally:
        conn.close()
